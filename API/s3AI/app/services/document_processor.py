import PyPDF2
import logging
from typing import Tuple
import tempfile
import os

# Optional: python-docx for DOCX support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

logger = logging.getLogger(__name__)

class DocumentProcessor:

    @staticmethod
    async def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text string
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"

            logger.info(f"Successfully extracted text from PDF: {file_path}")
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    async def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text string
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. DOCX support is unavailable. Install with: pip install python-docx")

        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            logger.info(f"Successfully extracted text from DOCX: {file_path}")
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    async def extract_text(file_path: str, file_type: str) -> str:
        """
        Extract text from any supported document type

        Args:
            file_path: Path to file
            file_type: Type of file (pdf, docx)

        Returns:
            Extracted text string
        """
        file_type = file_type.lower()

        if file_type == 'pdf':
            return await DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            return await DocumentProcessor.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    async def validate_file(file_path: str, file_type: str, max_size: int) -> Tuple[bool, str]:
        """
        Validate uploaded file

        Args:
            file_path: Path to file
            file_type: Type of file
            max_size: Maximum file size in bytes

        Returns:
            Tuple of (is_valid, message)
        """
        try:
            # Check file exists
            if not os.path.exists(file_path):
                return False, "File does not exist"

            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > max_size:
                return False, f"File size exceeds maximum allowed size ({max_size} bytes)"

            # Check file type
            supported_types = ['pdf']
            if DOCX_AVAILABLE:
                supported_types.append('docx')

            if file_type.lower() not in supported_types:
                return False, f"Unsupported file type. Supported types: {', '.join(supported_types)}"

            # Try to read file to ensure it's not corrupted
            if file_type.lower() == 'pdf':
                with open(file_path, 'rb') as f:
                    PyPDF2.PdfReader(f)
            elif file_type.lower() == 'docx':
                if not DOCX_AVAILABLE:
                    return False, "DOCX support is not available (python-docx not installed)"
                Document(file_path)

            return True, "File validation successful"

        except Exception as e:
            logger.error(f"Error validating file: {str(e)}")
            return False, f"File validation failed: {str(e)}"

    @staticmethod
    async def extract_metadata(file_path: str, file_type: str) -> dict:
        """
        Extract metadata from document

        Args:
            file_path: Path to file
            file_type: Type of file

        Returns:
            dict with metadata
        """
        try:
            file_size = os.path.getsize(file_path)
            metadata = {
                "file_size": file_size,
                "file_type": file_type
            }

            if file_type.lower() == 'pdf':
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    metadata["num_pages"] = len(pdf_reader.pages)

            elif file_type.lower() == 'docx':
                if not DOCX_AVAILABLE:
                    metadata["warning"] = "DOCX support not available (python-docx not installed)"
                    return metadata
                doc = Document(file_path)
                metadata["num_paragraphs"] = len(doc.paragraphs)
                metadata["num_tables"] = len(doc.tables)

            return metadata

        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {"error": str(e)}
