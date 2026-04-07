import boto3
import io
import logging
from typing import Optional
from app.config import settings

# PDF extraction libraries
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# DOCX extraction library
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)


class S3DocumentService:
    """Service for downloading and extracting text from documents in S3"""
    
    def __init__(self):
        try:
            self.s3_client = boto3.client(
                's3',
                region_name=settings.aws_region,
                aws_access_key_id=settings.aws_access_key_id,
                aws_secret_access_key=settings.aws_secret_access_key
            )
            self.bucket = settings.aws_s3_bucket
            logger.info(f"S3DocumentService initialized for bucket: {self.bucket}")
        except Exception as e:
            logger.error(f"Failed to initialize S3 client: {str(e)}")
            raise

    def download_document(self, s3_key: str) -> bytes:
        """
        Download a document from S3
        
        Args:
            s3_key: S3 key/path to the document
            
        Returns:
            Document content as bytes
            
        Raises:
            Exception: If download fails
        """
        try:
            logger.info(f"Downloading document from S3: {s3_key}")
            response = self.s3_client.get_object(Bucket=self.bucket, Key=s3_key)
            content = response['Body'].read()
            logger.debug(f"Downloaded {len(content)} bytes from S3: {s3_key}")
            return content
        except Exception as e:
            logger.error(f"Error downloading document from S3 ({s3_key}): {str(e)}")
            raise

    def extract_text_from_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extract text from PDF file
        
        Args:
            pdf_bytes: PDF file content as bytes
            
        Returns:
            Extracted text
            
        Raises:
            Exception: If extraction fails
        """
        try:
            logger.debug("Extracting text from PDF using available libraries")
            
            # Try pdfplumber first (more reliable for complex PDFs)
            if HAS_PDFPLUMBER:
                try:
                    pdf_file = io.BytesIO(pdf_bytes)
                    with pdfplumber.open(pdf_file) as pdf:
                        text = ""
                        for page_num, page in enumerate(pdf.pages):
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                            logger.debug(f"Extracted text from PDF page {page_num + 1}")
                    
                    logger.info(f"PDF extraction complete: {len(text)} characters extracted")
                    return text.strip()
                except Exception as e:
                    logger.warning(f"pdfplumber extraction failed: {str(e)}, trying PyPDF2")
            
            # Fallback to PyPDF2
            if HAS_PYPDF2:
                pdf_file = io.BytesIO(pdf_bytes)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    logger.debug(f"Extracted text from PDF page {page_num + 1}")
                
                logger.info(f"PDF extraction complete: {len(text)} characters extracted")
                return text.strip()
            
            # No library available
            raise RuntimeError(
                "No PDF extraction library available. "
                "Install PyPDF2 or pdfplumber: pip install PyPDF2 pdfplumber"
            )
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise

    def extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """
        Extract text from DOCX file
        
        Args:
            docx_bytes: DOCX file content as bytes
            
        Returns:
            Extracted text
            
        Raises:
            Exception: If extraction fails
        """
        try:
            if not HAS_DOCX:
                raise RuntimeError(
                    "python-docx library not available. "
                    "Install it: pip install python-docx"
                )
            
            logger.debug("Extracting text from DOCX")
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)
            
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            logger.info(f"DOCX extraction complete: {len(text)} characters extracted")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise

    def extract_text_from_document(self, file_bytes: bytes, file_type: str) -> str:
        """
        Extract text from a document based on file type
        
        Args:
            file_bytes: Document content as bytes
            file_type: File type ('pdf' or 'docx')
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If file type is not supported
            Exception: If extraction fails
        """
        file_type_lower = file_type.lower().strip()
        
        if file_type_lower == "pdf":
            return self.extract_text_from_pdf(file_bytes)
        elif file_type_lower == "docx":
            return self.extract_text_from_docx(file_bytes)
        else:
            raise ValueError(
                f"Unsupported file type: {file_type}. "
                f"Supported types: pdf, docx"
            )

    def get_document_text(self, s3_key: str, file_type: str) -> str:
        """
        Download and extract text from a document in S3
        
        Args:
            s3_key: S3 key/path to the document
            file_type: File type ('pdf' or 'docx')
            
        Returns:
            Extracted text
            
        Raises:
            Exception: If download or extraction fails
        """
        try:
            logger.info(f"Fetching document from S3: {s3_key} ({file_type})")
            document_bytes = self.download_document(s3_key)
            text = self.extract_text_from_document(document_bytes, file_type)
            logger.info(f"Successfully extracted {len(text)} characters from {s3_key}")
            return text
        except Exception as e:
            logger.error(f"Error getting document text from S3: {str(e)}")
            raise
